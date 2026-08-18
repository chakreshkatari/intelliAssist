"""
document_loader.py
-------------------
Handles text extraction from PDF, TXT, and DOCX files, plus preprocessing
and chunking of the extracted text so it can be embedded and indexed.
"""

from __future__ import annotations

import os
import re
import uuid
from dataclasses import dataclass, field
from typing import List

from pypdf import PdfReader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass
class Chunk:
    """A single chunk of text extracted from a source document."""
    chunk_id: str
    doc_name: str
    chunk_index: int
    text: str
    metadata: dict = field(default_factory=dict)


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
    except Exception as exc:
        raise ValueError(f"Could not open '{os.path.basename(file_path)}' as a PDF: {exc}") from exc

    pages = []
    for page_num, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # a single malformed page shouldn't sink the whole document
            print(f"[document_loader] Skipping unreadable page {page_num} in "
                  f"'{os.path.basename(file_path)}': {exc}")
            page_text = ""
        pages.append(page_text)

    if not any(p.strip() for p in pages):
        raise ValueError(
            f"No extractable text found in '{os.path.basename(file_path)}'. "
            f"It may be a scanned/image-only PDF, which this app cannot read."
        )
    return "\n".join(pages)


def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except OSError as exc:
        raise ValueError(f"Could not read '{os.path.basename(file_path)}': {exc}") from exc


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
    except Exception as exc:
        raise ValueError(f"Could not open '{os.path.basename(file_path)}' as a DOCX file: {exc}") from exc

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also pull text out of any tables in the document
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    if not parts:
        raise ValueError(f"No extractable text found in '{os.path.basename(file_path)}'.")
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )


# --------------------------------------------------------------------------
# Preprocessing
# --------------------------------------------------------------------------

def normalize_whitespace(text: str) -> str:
    """
    Collapse any run of whitespace (including embedded newlines) into a
    single space. Used as a final defensive pass on any text that will be
    displayed directly in the UI (chat answers, citation previews, search
    result snippets) so a stray newline can never render as a broken,
    vertical-looking line in Streamlit's markdown output.
    """
    return re.sub(r"\s+", " ", text).strip()


def clean_text(text: str) -> str:
    """
    Normalise whitespace and strip stray control characters.

    Both PDF and DOCX extraction commonly emit one newline per physical
    line of the source document rather than one newline per paragraph --
    a sentence that happened to wrap onto a new line in the original file
    comes back as two lines joined by a single '\\n', with no period in
    between. Left alone, that stray newline ends up embedded inside a
    single chunk/sentence and Streamlit's markdown renderer displays it as
    a hard line break, producing exactly the broken, vertical-looking text
    fragments this fixes (e.g. "...and\\npreprocess\\ntext..." instead of
    "...and preprocess text...").

    To fix this at the source: a run of 2+ newlines is treated as an
    intentional paragraph break and preserved as such, while a *single*
    newline (almost always just a mid-sentence line wrap from the source
    file) is collapsed into a plain space.
    """
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)          # normalise Windows/old-Mac line endings
    text = re.sub(r"[ \t]+", " ", text)            # collapse runs of spaces/tabs

    PARA_BREAK = "\u2029"                            # temporary paragraph-break marker
    text = re.sub(r"\n{2,}", PARA_BREAK, text)      # 2+ newlines = real paragraph break
    text = text.replace("\n", " ")                   # single newline = mid-sentence wrap -> space
    text = text.replace(PARA_BREAK, "\n\n")

    text = re.sub(r"[ \t]+", " ", text)             # collapse any double spaces just created
    paragraphs = [p.strip() for p in text.split("\n\n")]
    text = "\n\n".join(p for p in paragraphs if p)
    return text.strip()


# --------------------------------------------------------------------------
# Chunking
# --------------------------------------------------------------------------

def chunk_text(
    text: str,
    doc_name: str,
    chunk_size: int = 800,
    chunk_overlap: int = 120,
) -> List[Chunk]:
    """
    Split text into overlapping chunks using a sliding window over whole
    SENTENCES (never a raw character slice), so a RAG system can retrieve
    reasonably self-contained pieces of context with no broken fragments at
    chunk boundaries.

    Overlap is built from as many complete trailing sentences of the
    previous chunk as fit within chunk_overlap characters. This is
    deliberately sentence-aware rather than character-aware: an earlier,
    simpler implementation took a fixed-size trailing character slice and
    only trimmed a partial leading *word* -- which fixed broken words (e.g.
    "BERT" cut into "BE"/"RT") but could still land the cut in the middle of
    a sentence, silently dropping that sentence's leading words (e.g. a
    chunk starting with "(Bidirectional Encoder Representations..." instead
    of "BERT (Bidirectional Encoder Representations...") even though the cut
    itself was technically at a clean word boundary. Working in whole
    sentences avoids this entire class of bug: overlap text is always either
    a complete sentence or nothing at all, never a fragment of one.
    """
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    chunks: List[Chunk] = []
    current_sentences: List[str] = []
    current_len = 0
    chunk_index = 0

    def flush() -> None:
        nonlocal chunk_index
        if not current_sentences:
            return
        chunks.append(
            Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_name=doc_name,
                chunk_index=chunk_index,
                text=" ".join(current_sentences).strip(),
            )
        )
        chunk_index += 1

    for sentence in sentences:
        sentence_len = len(sentence) + 1  # +1 for the joining space
        fits = current_len + sentence_len <= chunk_size
        if fits or not current_sentences:
            # Always accept at least one sentence per chunk (even if it alone
            # exceeds chunk_size) so a single very long sentence can't create
            # an infinite loop of empty chunks.
            current_sentences.append(sentence)
            current_len += sentence_len
        else:
            flush()
            # Build the overlap seed from as many WHOLE trailing sentences
            # of the chunk just flushed as fit within chunk_overlap chars.
            overlap_sentences: List[str] = []
            overlap_len = 0
            for s in reversed(current_sentences):
                extra = len(s) + 1
                if overlap_len + extra > chunk_overlap:
                    break
                overlap_sentences.insert(0, s)
                overlap_len += extra
            current_sentences = overlap_sentences + [sentence]
            current_len = sum(len(s) + 1 for s in current_sentences)

    flush()
    return chunks


def load_and_chunk(file_path: str, chunk_size: int = 800, chunk_overlap: int = 120):
    """
    Full pipeline: extract -> clean -> chunk, for a single file on disk.
    Returns a tuple of (chunks, full_cleaned_text). full_cleaned_text is the
    de-duplicated, non-overlapping full document text -- used for whole
    document summarisation, since the chunk list itself contains
    intentionally overlapping text that would duplicate sentences if
    naively concatenated back together.
    """
    doc_name = os.path.basename(file_path)
    raw_text = extract_text(file_path)
    cleaned = clean_text(raw_text)
    if not cleaned:
        return [], ""
    chunks = chunk_text(cleaned, doc_name, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return chunks, cleaned
