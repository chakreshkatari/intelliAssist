"""
document_loader.py
-------------------
Handles text extraction from PDF, TXT, and DOCX files, plus preprocessing
and chunking of the extracted text so it can be embedded and indexed.
"""

from __future__ import annotations

import os
import re
import uuid
from dataclasses import dataclass, field
from typing import List

from pypdf import PdfReader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


@dataclass
class Chunk:
    """A single chunk of text extracted from a source document."""
    chunk_id: str
    doc_name: str
    chunk_index: int
    text: str
    metadata: dict = field(default_factory=dict)


# --------------------------------------------------------------------------
# Text extraction
# --------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)
    return "\n".join(pages)


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also pull text out of any tables in the document
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )


# --------------------------------------------------------------------------
# Preprocessing
# --------------------------------------------------------------------------

def clean_text(text: str) -> str:
    """Normalise whitespace and strip stray control characters."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


# --------------------------------------------------------------------------
# Chunking
# --------------------------------------------------------------------------

def chunk_text(
    text: str,
    doc_name: str,
    chunk_size: int = 800,
    chunk_overlap: int = 120,
) -> List[Chunk]:
    """
    Split text into overlapping chunks using a simple sliding window over
    sentences, so a RAG system can retrieve reasonably self-contained pieces
    of context. chunk_size / chunk_overlap are measured in characters.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks: List[Chunk] = []
    current = ""
    chunk_index = 0

    for sentence in sentences:
        if not sentence.strip():
            continue
        if len(current) + len(sentence) + 1 <= chunk_size:
            current = f"{current} {sentence}".strip()
        else:
            if current:
                chunks.append(
                    Chunk(
                        chunk_id=str(uuid.uuid4()),
                        doc_name=doc_name,
                        chunk_index=chunk_index,
                        text=current.strip(),
                    )
                )
                chunk_index += 1
                # start next chunk with overlap from the tail of the previous one
                overlap_text = current[-chunk_overlap:] if chunk_overlap else ""
                current = f"{overlap_text} {sentence}".strip()
            else:
                current = sentence

    if current.strip():
        chunks.append(
            Chunk(
                chunk_id=str(uuid.uuid4()),
                doc_name=doc_name,
                chunk_index=chunk_index,
                text=current.strip(),
            )
        )

    return chunks


def load_and_chunk(file_path: str, chunk_size: int = 800, chunk_overlap: int = 120):
    """
    Full pipeline: extract -> clean -> chunk, for a single file on disk.
    Returns a tuple of (chunks, full_cleaned_text). full_cleaned_text is the
    de-duplicated, non-overlapping full document text -- used for whole
    document summarisation, since the chunk list itself contains
    intentionally overlapping text that would duplicate sentences if
    naively concatenated back together.
    """
    doc_name = os.path.basename(file_path)
    raw_text = extract_text(file_path)
    cleaned = clean_text(raw_text)
    if not cleaned:
        return [], ""
    chunks = chunk_text(cleaned, doc_name, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return chunks, cleaned
